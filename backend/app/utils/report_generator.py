import io
from docx import Document
from docx.shared import Inches
import jinja2
from datetime import datetime
from typing import Dict, Any

def generate_docx_report(session, test, report_type="client"):
    """Generate DOCX report from test session."""
    doc = Document()
    doc.add_heading(test.title, 0)
    doc.add_paragraph(f"Client: {session.client_name}")
    doc.add_paragraph(f"Completed: {session.completed_at.strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading("Answers", level=1)
    # get questions mapping
    questions = {q["id"]: q for q in test.config["questions"]}
    for qid, answer in session.answers.items():
        q = questions.get(qid, {"text": qid})
        doc.add_paragraph(f"{q['text']}: {answer}")
    if session.metrics:
        doc.add_heading("Calculated Metrics", level=1)
        for name, value in session.metrics.items():
            doc.add_paragraph(f"{name}: {value}")
    # save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

def generate_html_report(session, test, report_type="client"):
    """Generate HTML report using Jinja2 template."""
    env = jinja2.Environment(loader=jinja2.FileSystemLoader("app/templates"))
    template = env.get_template(f"{report_type}_report.html")
    questions = {q["id"]: q for q in test.config["questions"]}
    answers_with_text = []
    for qid, answer in session.answers.items():
        q = questions.get(qid, {"text": qid})
        answers_with_text.append({"question": q["text"], "answer": answer})
    html = template.render(
        title=test.title,
        client_name=session.client_name,
        completed_at=session.completed_at,
        answers=answers_with_text,
        metrics=session.metrics or {}
    )
    return html